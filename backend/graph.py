from typing_extensions import TypedDict
from llm import model
from prompt import drafting_prompt_template, drafting_chat_prompt
from vector_rag import retriever
from langgraph.graph import StateGraph, START, END
import os
import docx
import json
import re
from transformers import pipeline
from docx.shared import Pt
from docx.shared import Inches # Import Inches for table column width
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Import alignment
from docx.shared import Pt # Import Pt for font size (basic attempt)
from docx.enum.style import WD_STYLE_TYPE # Import style type
from docx.shared import RGBColor # Import RGBColor for font color
from docx.enum.table import WD_TABLE_ALIGNMENT # Import table alignment

# Helper function to set all text in a paragraph to black
def set_text_color_to_black(paragraph):
    """Ensure all text in a paragraph is black."""
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    # Font size and color will be handled by modifying document styles

# using the new ComplianceAgent from compliance_checker.py.
from compliance_checker import ComplianceAgent

# Initialize the toxicity classifier.
toxicity_classifier = pipeline("text-classification", model="unitary/unbiased-toxic-roberta")

# Define our state
class State(TypedDict, total=False):
    flow: str
    previous_sow: str
    query_map: dict
    user_query: str
    additional_context: str
    sow: str            # SOW as a JSON string (or raw text) produced by the drafting agent.
    validated_sow: dict # Parsed and validated SOW data.
    formatted_sow: str  # Filename of the generated DOCX.
    compliance_results: dict  # Results from compliance analysis.
    feedback: str
    error: str
    retryCount: int
    doc_file_path: str

def get_relevant_context(state: State):
    context = retriever.invoke(state['user_query'])
    return { 'additional_context': context, 'retryCount': 0 }

def extract_raw_json(response_text):
    """
    Extracts and parses raw JSON from an AI response wrapped in ```json ... ```
    or returns the raw JSON directly if no wrapping exists.
    """
    # Pattern to capture content between ```json ... ``` or just ```
    pattern = r"```(?:json)?\s*([\s\S]*?)```"

    match = re.search(pattern, response_text.strip())
    if match:
        raw_json_str = match.group(1)
    else:
        raw_json_str = response_text.strip()

    try:
        return json.loads(raw_json_str)
    except json.JSONDecodeError as e:
        print(f"❌ JSON parsing failed: {e}")
        return None


def drafting_agent(state: State):
    if state.get('error'):
        previous_content = state.get('sow', '')
        instruction = (f"Below is the previously generated content: {previous_content} "
                       f"The following errors were detected: {state['error']}. "
                       f"Please revise the content accordingly.")
    else:
        instruction = ""

    if state.get('flow') == 'chat':
        prompt = drafting_chat_prompt.invoke({
            "user_query": state['user_query'],
            "previous_sow": state['previous_sow'],
            "feedback": instruction,
        })
        response = model.invoke(prompt)
        return { 'sow': response.content }
    else: 
        prompt = drafting_prompt_template.invoke({
            "query": state['user_query'],
            "additional_context": state['additional_context'],
            "feedback": instruction,
            **state['query_map']
        })
        response = model.invoke(prompt)
        return { 'sow': response.content }

def compliance_agent(state: State):
    # Use the  ComplianceAgent to analyze the SOW.
    agent = ComplianceAgent()
    try:
        # Attempt to parse the SOW as JSON
        sow_data = json.loads(state['sow'])
    except Exception:
        # If parsing fails, treat the entire text as the content to check.
        sow_data = {"sow_text": state['sow']}
    
    # Generate the compliance report using the new agent.
    report = agent.generate_report(sow_data)
    state['compliance_results'] = report
    
    # If any compliance issues are detected, build a brief error message and set it on the state.
    if (report["compliance_score"] < 80 or
        report["missing_fields"] or
        report["structural_issues"] or
        report["content_issues"] or
        report["language_issues"]):
        error_message = "Compliance issues detected: "
        if report["missing_fields"]:
            error_message += f"Missing fields: {', '.join(report['missing_fields'])}. "
        if report["structural_issues"]:
            error_message += f"Structural issues: {', '.join(report['structural_issues'])}. "
        if report["content_issues"]:
            error_message += f"Content issues: {', '.join(report['content_issues'])}. "
        if report["language_issues"]:
            error_message += f"Language issues: {', '.join(report['language_issues'])}. "
        error_message += f"Risk Level: {report['risk_level']}"
        state['error'] = error_message
    return state

def validate_text(text, threshold=0.75):
    try:
        result = toxicity_classifier(text)[0]
    except Exception as e:
        return text, None
    
    toxic_labels = [
        "toxicity", "severe_toxicity", "obscene", "threat",
        "insult", "identity_attack", "sexual_explicit"
    ]
    if result['label'] in toxic_labels and result['score'] > threshold:
        error_msg = (f"[⚠ TOXIC CONTENT DETECTED] Text validation failed: {text}. "
                     f"Reason: {result['label']} with score {round(result['score']*100, 2)}%")
        print(error_msg)
        return text, error_msg
    return text, None

def validate_sow_data(sow_data):
    validated_data = {}
    errors = {}
    for key, value in sow_data.items():
        if isinstance(value, dict):
            validated_subsection = {}
            subsection_errors = {}
            for subkey, subvalue in value.items():
                valid_text, error = validate_text(subvalue)
                validated_subsection[subkey] = valid_text
                if error:
                    subsection_errors[subkey] = error
            validated_data[key] = validated_subsection
            if subsection_errors:
                errors[key] = subsection_errors
        else:
            valid_text, error = validate_text(value)
            validated_data[key] = valid_text
            if error:
                errors[key] = error
    return validated_data, errors

def extract_json_from_sow(raw_sow: str) -> dict:
    extraction_prompt = (
        '''
        Extract the following fields from the given Statement of Work into a JSON object with these keys: 
        "Project Name", "End Date", "Confidentiality", "Intellectual Property", "Termination", "Project Title", "Start Date", "End Date", "Project Name", "SOW Effective Date","Company Information", "Client", "Agreement Date",
"Client Contact", "Contact", "Services Description", "Deliverables",
"Milestones", "Acceptance", "Personnel and Locations", "Representatives",
"Client Representatives", "Contractor Resources", "Terms & Conditions", "Fees", "Expenses",
"Taxes", "Conversion", "Limitation of Liability", "Service Level Agreement", "Assumptions", "Scope of Work",
"Change Process", "Payment Terms", "Timeline", "Company Name", "Client Name",
''' + raw_sow +
        "\n\nOutput the result as a valid JSON and do not format just return pure json."
    )
    response = model.invoke(extraction_prompt)
    try:
        sow_data = json.loads(response.content)
        return sow_data
    except Exception as e:
        raise ValueError("Failed to extract JSON from SOW content: " + str(e))

def validation_agent(state: State):
    try:
        try:
            sow_data = json.loads(state['sow'])
        except Exception as parse_error:
            sow_data = extract_json_from_sow(state['sow'])

        validated_data, errors = validate_sow_data(sow_data)
        if errors:
            state['error'] = json.dumps(errors)
            return { 'feedback': 'REJECTED', 'retryCount': state['retryCount'] + 1 }
        else:
            state['validated_sow'] = validated_data
            state.pop('error', None)
            return { 'feedback': 'ACCEPTED', 'validated_sow': validated_data }
    except Exception as e:
        state['error'] = str(e)
        return { 'feedback': 'REJECTED', 'retryCount': state['retryCount'] + 1 }

def generate_sow(sow_data, output_filename="Generated_SOW_final.docx"):
    """
    Generates a Statement of Work (SOW) DOCX file from a dictionary of data.

    Args:
        sow_data (dict): A dictionary containing the data for the SOW.
        output_filename (str, optional): The name of the output DOCX file.
            Defaults to "Generated_SOW_final.docx".
    """
    doc = docx.Document()
    # --- Styles ---
    # Define a consistent font
    font_name = 'Arial'
    doc.styles['Normal'].font.name = font_name
    doc.styles['Normal'].font.size = Pt(8)

    # Define heading styles with Arial and adjusted sizes
    heading_1 = doc.styles['Heading 1']
    heading_1.font.name = font_name
    heading_1.font.size = Pt(12)
    heading_1.font.color.rgb = RGBColor(0, 0, 0)   # Slightly larger than default

    heading_2 = doc.styles['Heading 2']
    heading_2.font.name = font_name
    heading_2.underline = True
    heading_2.font.size = Pt(11)
    heading_2.font.color.rgb = RGBColor(0, 0, 0) # Slightly larger

    def get_data(key, default="[Data Not Provided]"):
        """Helper to get nested data from the sow_data dictionary."""
        keys = key.split('.')
        val = sow_data
        for k in keys:
            if isinstance(val, dict) and k in val:
                val = val[k]
            else:
                return default
        return val if val is not None else default

    # --- Main SOW Document ---
    # --- Main SOW Document ---
# Add the main title "STATEMENT OF WORK" with a reduced font size
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("STATEMENT OF WORK")
    title_run.bold = True
    title_run.font.size = Pt(14)  # Set the font size (e.g., 14pt)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the Project Name under the main heading and center-align it
   # Add the Project Name under the main heading and center-align it
    project_name_paragraph = doc.add_paragraph()
    project_name_run = project_name_paragraph.add_run(get_data("Project Name", "[Project Name]"))
    project_name_run.bold = True  # Make the Project Name bold
    project_name_run.font.size = Pt(12)  # Increase the font size (e.g., 12pt)
    project_name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    project_name_paragraph.paragraph_format.space_before = Pt(6)  # Add some space before the project name
    project_name_paragraph.paragraph_format.space_after = Pt(12)    # Add some space after the project name

    # Add introductory text
    intro_p1 = (
        f"This Statement of Work (“SOW”) is entered into as of {get_data('SOW Effective Date')} "
        f"(the “SOW Effective Date”) by and between {get_data('Insight Global Full Name', 'Evergreen, a professional services division of Insight Global, LLC,')} (“Insight Global”) "
        f"and {get_data('Client Full Name')} (“Client”) under the provisions of that certain Master Services Agreement, "
        f"dated as of {get_data('Agreement Date')}, by and between Insight Global and Client (the “Agreement”)."
    )
    paragraph = doc.add_paragraph(intro_p1)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    doc.add_paragraph(
        "This SOW, and all attachments, exhibits and schedules hereto, are subject to the terms and provisions of the Agreement."
    ).paragraph_format.space_after = Pt(0)
    doc.add_paragraph(
        "In the event of any conflict between this SOW and the Agreement, the terms and conditions of this SOW will govern and prevail."
    ).paragraph_format.space_after = Pt(0)
    doc.add_paragraph(
        "Capitalized terms used and not otherwise defined herein shall have the meanings ascribed thereto in the Agreement."
    ).paragraph_format.space_after = Pt(0)

    # --- Contact Information ---
    doc.add_heading("Contact Information.", level=1)

    contact_table = doc.add_table(rows=7, cols=2)
    contact_table.style = 'Table Grid'
    contact_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    contact_table.allow_autofit = False

    # Set column widths
    contact_table.cell(0, 0).width = Inches(2.0)
    contact_table.cell(0, 1).width = Inches(4.0)

    contact_table.cell(0, 0).text = "Client Contact"
    contact_table.cell(0, 1).text = "Insight Global Contact"

    contact_fields_map = {
        "Name:": "Name", "Title:": "Title", "Address:": "Address",
        "Phone:": "Phone", "Mobile:": "Mobile", "Email:": "Email"
    }
    for i, (label, key) in enumerate(contact_fields_map.items()):
        row_idx = i + 1
        cell_0 = contact_table.cell(row_idx, 0)
        cell_1 = contact_table.cell(row_idx, 1)
        cell_0.text = f"{label} {get_data(f'Client Contact.{key}')} "
        cell_1.text = f"{label} {get_data(f'Insight Global Contact.{key}')}"
    section_order = [
    ("Definitions.", "Definitions"),
    ("Services Description.", "Services Description"),
    ("Deliverables.", "Deliverables"),
    ("Milestones, Timelines, and Delivery Dates.", "Milestones Timelines and Delivery Dates"),
    ("Acceptance.", "Acceptance Criteria"),
    ("Personnel and Locations.", "Location of Services"),
    ("Term.", "Term"),
    ("Fees.", "Fees Text"),
    ("Expenses.", "Expenses Text"),
    ("Taxes.", "Taxes Text"),
    ("Conversion.", "Conversion Text"),
    ("Limitation of Liability.", "Limitation of Liability Text"),
    ("Service Level Agreement.", "Service Level Agreement Text"),
    ("Assumptions.", "Assumptions.General Assumptions Text"),
    ("Change Process.", "Change Process Text"),
    ("Client Responsibilities.", "Assumptions.Client Responsibilities"),
    ("Overtime.", "Assumptions.Overtime"),
    ("Dependencies.", "Assumptions.Dependencies"),
    ("Equipment, Software, and Access.", "Assumptions.Equipment Software and Access"),
    ("Current State Analysis.", "Current State Analysis"),
    ("Gap Analysis.", "Gap Analysis"),
]
    # --- Main SOW Sections ---
    # Helper function to add numbered subpoints
    def add_numbered_subpoint(doc, main_number, sub_number, text):
        subpoint_paragraph = doc.add_paragraph()
        subpoint_paragraph.add_run(f"{main_number}.{sub_number} ").bold = True
        subpoint_paragraph.add_run(text)
        subpoint_paragraph.paragraph_format.space_before = Pt(0)
        subpoint_paragraph.paragraph_format.space_after = Pt(0)

    main_number = 1  # Start with the first main section
    for heading_text, data_key in section_order:
        doc.add_heading(f"{main_number}. {heading_text}", level=1)

        if heading_text == "Acceptance.":
            add_numbered_subpoint(doc, main_number, 1, get_data(data_key))
            add_numbered_subpoint(doc, main_number, 2, "Client will sign a weekly timesheet for each Insight Global resource.")
            add_numbered_subpoint(doc, main_number, 3, "By signing this timesheet, Client provides acceptance for all work performed towards the scope of the entire project during these hours.")
        elif heading_text == "Personnel and Locations.":
            add_numbered_subpoint(doc, main_number, 1, f"Location of Services: {get_data('Location of Services', 'The Services will be provided at the following location(s): [Specify Location(s)]')}")
            add_numbered_subpoint(doc, main_number, 2, f"Insight Global Representatives: {get_data('Insight Global Representatives Text', 'Insight Global shall be represented by the following personnel: [Specify IG Reps]')}")
            add_numbered_subpoint(doc, main_number, 3, f"Client Representatives: {get_data('Client Representatives Text', 'Client shall make the following representatives available, as reasonably required for Insight Global to perform the Services and Complete the Deliverables in a timely manner: [Specify Client Reps]')}")
        elif heading_text == "Fees.":
            add_numbered_subpoint(doc, main_number, 1, get_data(data_key))
            add_numbered_subpoint(doc, main_number, 2, "Should any Insight Global resource work over forty (40) hours per week and Insight Global is legally responsible to provide overtime compensation to the resource, overtime will be billed at a rate of one and a half (1.5) times the standard bill rate.")
        elif heading_text == "Assumptions.":
            doc.add_paragraph(get_data("Assumptions.General Assumptions Text", "The following additional assumptions and terms and conditions apply to this SOW:"))
            assumptions_list = {
                "Client Responsibilities.": "Assumptions.Client Responsibilities",
                "Overtime.": "Assumptions.Overtime",
                "Dependencies.": "Assumptions.Dependencies",
                "Equipment, Software, and Access.": "Assumptions.Equipment Software and Access"
            }
            sub_number = 1
            for sub_heading, sub_key in assumptions_list.items():
                add_numbered_subpoint(doc, main_number, sub_number, f"{sub_heading} {get_data(sub_key)}")
                sub_number += 1
        elif data_key:
            add_numbered_subpoint(doc, main_number, 1, get_data(data_key))
        else:
            doc.add_paragraph("[Content for this section is template-based or handled elsewhere]")

        main_number += 1  # Increment the main section number

    # IN WITNESS WHEREOF (SOW)
    doc.add_heading("IN WITNESS WHEREOF", level=1)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        "the signatories below hereto have caused this SOW to be executed by their duly authorized representatives effective as of the effective date of this SOW.")
    # doc.add_paragraph()

    sig_sow_table = doc.add_table(rows=5, cols=2)  # Header + 4 fields
    sig_sow_table.style = 'Table Grid'
    sig_sow_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    sig_sow_table.allow_autofit = False

    # Set width
    sig_sow_table.cell(0, 0).width = Inches(3.0)
    sig_sow_table.cell(0, 1).width = Inches(3.0)

    sig_sow_table.cell(0, 0).paragraphs[0].add_run("INSIGHT GLOBAL, LLC:").bold = True
    sig_sow_table.cell(0, 1).paragraphs[0].add_run("CLIENT:").bold = True

    sig_fields = ["By:", "Name:", "Title:", "Date:"]
    for i, field in enumerate(sig_fields):
        cell_0 = sig_sow_table.cell(i + 1, 0)
        cell_1 = sig_sow_table.cell(i + 1, 1)
        cell_0.text = field
        cell_1.text = field
        if field == "Name:":
            cell_0.text = f"Name: {get_data('Insight Global Signatory Name', '________________________________')}"
            cell_1.text = f"Name: {get_data('Client Signatory Name', '________________________________')}"
        elif field == "Title:":
            cell_0.text = f"Title: {get_data('Insight Global Signatory Title', '________________________________')}"
            cell_1.text = f"Title: {get_data('Client Signatory Title', '________________________________')}"
        elif field == "Date:":
            cell_0.text = f"Date: {get_data('Insight Global Signature Date', '__________________________________')}"
            cell_1.text = f"Date: {get_data('Client Signature Date', '__________________________________')}"
        else:  # By:
            cell_0.text = f"By: _________________________"
            cell_1.text = f"By: _________________________"

        # Make "By:", "Name:" etc bold
        # for paragraph in cell_0.paragraphs:
        #     for run in paragraph.runs:
        #         if field in run.text:
        #             run.font.bold = True
        # for paragraph in cell_1.paragraphs:
        #     for run in paragraph.runs:
        #         if field in run.text:
        #             run.font.bold = True

    # --- ATTACHMENT A: CHANGE ORDER ---
    doc.add_page_break()
    doc.add_heading("ATTACHMENT A:", level=1)
    doc.add_heading("CHANGE ORDER NO. ___", level=2)  # Number can be dynamic if needed
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create the table for "Change requested in" and "Affected Section numbers of SOW"
    attachment_table = doc.add_table(rows=2, cols=2)
    attachment_table.style = 'Table Grid'
    attachment_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    attachment_table.allow_autofit = False

    # Set column widths
    attachment_table.columns[0].width = Inches(4.0)  # Wider column for checkboxes
    attachment_table.columns[1].width = Inches(2.5)  # Narrower column for affected sections

    # Populate the first row
    attachment_table.cell(0, 0).text = "Change requested in: (Check all that apply)"
    checkboxes_text = (
        f"{'☐' if not get_data('CO Change Requested Services', False) else '☒'} Services\n"
        f"{'☐' if not get_data('CO Change Requested Deliverables', False) else '☒'} Deliverables\n"
        f"{'☐' if not get_data('CO Change Requested Timeline', False) else '☒'} Timeline\n"
        f"{'☐' if not get_data('CO Change Requested Fees', False) else '☒'} Fees\n"
        f"{'☐' if not get_data('CO Change Requested Other', False) else '☒'} Other"
    )
    attachment_table.cell(0, 0).text += f"\n{checkboxes_text}"

    attachment_table.cell(0, 1).text = "Affected Section numbers of SOW:\n" + get_data("CO Affected SOW Sections", "•")

    # Merge the second row for "Change Description and Reason"
    attachment_table.cell(1, 0).merge(attachment_table.cell(1, 1))
    attachment_table.cell(1, 0).text = "Change Description and Reason:\n" + get_data("CO Change Description and Reason", "[Specify description and reason]")

    # Format the table cells
    for row in attachment_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
        doc.add_paragraph("Changes to SOW: The Parties hereto agree as follows:")
        doc.add_paragraph(
            f"Section {get_data('CO Section Number', '#')}: {get_data('CO Section Name', 'Section Name')} "
            f"is hereby deleted in its entirety and replaced with the following:"
        )
        doc.add_paragraph(
            f"{get_data('CO Section Number', '#')}. {get_data('CO Section Name', 'Section Name')}. {get_data('CO New Section Terms', 'New section terms')}")
        # doc.add_paragraph()

        doc.add_paragraph("[Remainder of page intentionally left blank]")
        doc.add_paragraph()


        doc.add_heading("IN WITNESS WHEREOF", level=2)  # Level 2 for CO
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(
            "the authorized representatives of the parties have executed this Change Order to become effective as of the above Change Order Effective Date.")
        # doc.add_paragraph()

        co_sig_table = doc.add_table(rows=4, cols=3)  # Labels in col 0, IG in col 1, Client in col 2
        co_sig_table.style = 'Table Grid'
        co_sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        co_sig_table.allow_autofit = False

        co_sig_table.cell(0, 0).width = Inches(1.5)
        co_sig_table.cell(0, 1).width = Inches(2.0)
        co_sig_table.cell(0, 2).width = Inches(2.0)
        co_sig_table.cell(0, 1).paragraphs[0].add_run("INSIGHT GLOBAL, LLC").bold = True
        co_sig_table.cell(0, 2).paragraphs[0].add_run("CLIENT").bold = True

        co_sig_fields = ["Signature", "Printed Name", "Title", "Date"]
        for i, field_label in enumerate(co_sig_fields):
            if field_label == "Signature":
                co_sig_table.cell(i, 0).text = ""
                co_sig_table.cell(i, 1).text = "Signature _________________________"
                co_sig_table.cell(i, 2).text = "Signature _________________________"
            else:
                co_sig_table.cell(i, 0).text = field_label
                ig_val_key = f"CO IG Signatory {field_label.replace(' ', '')}"
                client_val_key = f"CO Client Signatory {field_label.replace(' ', '')}"
                ig_val = get_data(ig_val_key, "________________________")
                client_val = get_data(client_val_key, "________________________")
                co_sig_table.cell(i, 1).text = f"{field_label}  {ig_val}"
                co_sig_table.cell(i, 2).text = f"{field_label}  {client_val}"

            # Make "Signature", "Printed Name" etc bold.
            if field_label != "Signature":
                for paragraph in co_sig_table.cell(i,0).paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

        # --- Saving the Document ---
        try:
            current_script_path = os.path.dirname(__file__)
        except NameError:
            current_script_path = os.getcwd()  # Fallback

        static_folder = os.path.join(current_script_path, 'static')
        if not os.path.exists(static_folder):
            os.makedirs(static_folder)

        output_path = os.path.join(static_folder, output_filename)
        doc.save(output_path)

        print(f"✅ SOW document generated: {output_path}")
        return  output_path # returning the path

def formatting_agent(state: State):
    output_file = generate_sow(state['validated_sow'])
    state['doc_file_path'] = output_file['fileName']
    state['formatted_sow'] = output_file['formatted_sow_md']
    return state

def agent_router(state: State):
    if state['retryCount'] > 10:
        return 'SUCCESS'
    # If any error exists (from compliance or validation), loop back to drafting.
    if state.get('error'):
        print(f"error: {state.get('error')}")
        return 'REJECTED'
    if state.get('feedback') == 'ACCEPTED':
        return 'SUCCESS'
    print(f"error: {state.get('error')}")
    return 'REJECTED'

# ---- Build the Graph ----
graph_builder = StateGraph(State)
graph_builder.add_node('get_relevant_context', get_relevant_context)
graph_builder.add_node('drafting_agent', drafting_agent)
graph_builder.add_node('compliance_agent', compliance_agent)
graph_builder.add_node('validation_agent', validation_agent)
graph_builder.add_node('formatting_agent', formatting_agent)

graph_builder.add_edge(START, 'get_relevant_context')
graph_builder.add_edge('get_relevant_context', 'drafting_agent')
graph_builder.add_edge('drafting_agent', 'compliance_agent')
graph_builder.add_edge('compliance_agent', 'validation_agent')

# Route based on validation and compliance feedback.
graph_builder.add_conditional_edges(
    "validation_agent",
    agent_router,
    {
        'SUCCESS': 'formatting_agent',
        'REJECTED': 'drafting_agent'
    }
)

graph_builder.add_edge('formatting_agent', END)

graph_agentor = graph_builder.compile()